#!/usr/bin/env python3
"""Convierte docs/reporte-tecnico.md a docs/reporte-tecnico.docx (Word).

No depende de pandoc -- usa python-docx directamente, con soporte para lo
que realmente usa el reporte: encabezados, tablas, bloques de código,
listas, negrita/cursiva/código inline, citas, y las 3 capturas de pantalla
insertadas donde el markdown las referencia.

Uso:
    python scripts/generate-report-docx.py
"""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parent.parent
MD_PATH = ROOT / "docs" / "reporte-tecnico.md"
DOCX_PATH = ROOT / "docs" / "reporte-tecnico.docx"
SCREENSHOTS_DIR = ROOT / "docs" / "screenshots"

CODE_FONT = "Consolas"
BODY_FONT = "Calibri"

INLINE_RE = re.compile(r"(\*\*.+?\*\*|`.+?`|\*[^*]+?\*|\[.+?\]\(.+?\))")
LINK_RE = re.compile(r"^\[(.+)\]\((.+)\)$")


def setup_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(8)

    for level, size in ((1, 18), (2, 14), (3, 12)):
        style = doc.styles[f"Heading {level}"]
        style.font.name = BODY_FONT
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
        style.paragraph_format.space_before = Pt(18 if level == 1 else 10)
        style.paragraph_format.space_after = Pt(8)

    if "CodeBlock" not in doc.styles:
        code_style = doc.styles.add_style("CodeBlock", 1)  # WD_STYLE_TYPE.PARAGRAPH
        code_style.font.name = CODE_FONT
        code_style.font.size = Pt(9)
        code_style.paragraph_format.space_before = Pt(4)
        code_style.paragraph_format.space_after = Pt(10)
        code_style.paragraph_format.left_indent = Cm(0.4)


def shade_paragraph(paragraph, hex_color: str) -> None:
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    pPr.append(shd)


def add_inline_runs(paragraph, text: str) -> None:
    for token in INLINE_RE.split(text):
        if not token:
            continue
        link_match = LINK_RE.match(token)
        if link_match:
            label = link_match.group(1).strip("`")
            run = paragraph.add_run(label)
            run.font.name = CODE_FONT if "`" in link_match.group(1) else BODY_FONT
            run.italic = True
            run.font.color.rgb = RGBColor(0x2E, 0x5B, 0x9E)
            continue
        if token.startswith("**") and token.endswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        elif token.startswith("`") and token.endswith("`"):
            run = paragraph.add_run(token[1:-1])
            run.font.name = CODE_FONT
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0xB0, 0x30, 0x60)
        elif token.startswith("*") and token.endswith("*"):
            run = paragraph.add_run(token[1:-1])
            run.italic = True
        else:
            paragraph.add_run(token)


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Light Grid Accent 1"
    for r, row in enumerate(rows):
        for c, cell_text in enumerate(row):
            cell = table.rows[r].cells[c]
            cell.paragraphs[0].text = ""
            add_inline_runs(cell.paragraphs[0], cell_text)
            if r == 0:
                for run in cell.paragraphs[0].runs:
                    run.bold = True
    doc.add_paragraph()


def add_image_if_referenced(doc: Document, line: str) -> bool:
    m = re.search(r"screenshots/([\w.\-]+\.png)", line)
    if not m:
        return False
    img_path = SCREENSHOTS_DIR / m.group(1)
    if not img_path.exists():
        return False
    doc.add_picture(str(img_path), width=Cm(16))
    caption = doc.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = caption.add_run(m.group(1))
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
    return True


def convert(md_text: str) -> Document:
    doc = Document()
    for section in doc.sections:
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
    setup_styles(doc)

    lines = md_text.splitlines()
    i = 0
    n = len(lines)
    in_frontmatter = False
    title_done = False

    while i < n:
        line = lines[i]

        # Frontmatter YAML
        if i == 0 and line.strip() == "---":
            in_frontmatter = True
            i += 1
            continue
        if in_frontmatter:
            if line.strip() == "---":
                in_frontmatter = False
            i += 1
            continue

        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        if stripped == "---":
            i += 1
            continue

        # Fenced code block
        if stripped.startswith("```"):
            i += 1
            code_lines = []
            while i < n and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            i += 1  # skip closing fence
            code_text = "\n".join(code_lines) if code_lines else " "
            p = doc.add_paragraph(style="CodeBlock")
            p.add_run(code_text)
            shade_paragraph(p, "F2F2F2")
            continue

        # Table
        if stripped.startswith("|"):
            table_rows = []
            while i < n and lines[i].strip().startswith("|"):
                row_line = lines[i].strip().strip("|")
                cells = [c.strip() for c in row_line.split("|")]
                if not re.match(r"^:?-+:?$", cells[0]):
                    table_rows.append(cells)
                i += 1
            add_table(doc, table_rows)
            continue

        # Blockquote
        if stripped.startswith(">"):
            quote_lines = []
            while i < n and lines[i].strip().startswith(">"):
                quote_lines.append(lines[i].strip().lstrip(">").strip())
                i += 1
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.8)
            run = p.add_run(" ".join(quote_lines))
            run.italic = True
            run.font.color.rgb = RGBColor(0x40, 0x40, 0x40)
            continue

        # Headings
        heading_match = re.match(r"^(#{1,3})\s+(.*)$", stripped)
        if heading_match:
            hashes, text = heading_match.groups()
            level = len(hashes)
            if level == 1 and not title_done:
                p = doc.add_paragraph()
                p.style = doc.styles["Title"] if "Title" in [s.name for s in doc.styles] else doc.styles["Heading 1"]
                add_inline_runs(p, text)
                title_done = True
            else:
                # '#' (level 1) es el título del documento (arriba); '##' es
                # el nivel principal de sección del reporte -> Heading 1 de
                # Word, '###' -> Heading 2.
                doc_level = max(1, level - 1)
                p = doc.add_heading(level=doc_level)
                add_inline_runs(p, text)
            add_image_if_referenced(doc, text)
            i += 1
            continue

        # Bullet list (con líneas de continuación indentadas bajo el mismo ítem)
        if stripped.startswith("- "):
            item_lines = [stripped[2:]]
            i += 1
            while i < n and lines[i].strip() and not lines[i].strip().startswith(
                ("#", "```", "|", ">", "- ", "---")
            ):
                item_lines.append(lines[i].strip())
                i += 1
            full_item = " ".join(item_lines)
            p = doc.add_paragraph(style="List Bullet")
            add_inline_runs(p, full_item)
            add_image_if_referenced(doc, full_item)
            continue

        # Regular paragraph (accumulate continuation lines until blank/structural)
        para_lines = [stripped]
        i += 1
        while i < n and lines[i].strip() and not lines[i].strip().startswith(("#", "```", "|", ">", "- ", "---")):
            para_lines.append(lines[i].strip())
            i += 1
        full_text = " ".join(para_lines)
        p = doc.add_paragraph()
        add_inline_runs(p, full_text)
        add_image_if_referenced(doc, full_text)

    return doc


def main() -> None:
    md_text = MD_PATH.read_text(encoding="utf-8")
    doc = convert(md_text)
    doc.save(str(DOCX_PATH))
    print(f"Generado: {DOCX_PATH}")


if __name__ == "__main__":
    main()
